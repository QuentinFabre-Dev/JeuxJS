#!/usr/bin/env python3
"""
Builds the sample contract used for demos, from a single source of truth.

Outputs:
  public/samples/contrat-prestation-services.docx   the file to import
  src/data/mockDocument.js                          demo-mode preview
  src/data/mockFindings.js                          demo-mode findings

The demo engine matches findings to sentences by exact text, so generating both
from the same source is the only way to keep them in step. Editing the contract
by hand in one place and not the other silently breaks the highlighting.

Requires: pip install python-docx
Run from the repository root: python3 tools/build-sample-contract.py
"""

import json
import os

# ── The contract ───────────────────────────────────────────────────────────
#
# Each entry is either a heading or a paragraph. A paragraph is a list of
# sentences: that is the unit the app reviews, so defects are planted at the
# sentence level.
#
# `defect` marks a sentence the review is expected to flag, with everything
# needed to build the matching demo finding.

CONTRACT = [
    ("heading", "CONTRAT DE PRESTATION DE SERVICES"),
    ("p", [
        "Entre les soussignés : NOVATEK Conseil SAS, société par actions simplifiée au capital de 50 000 euros, immatriculée au RCS de Paris sous le numéro 892 451 337, dont le siège social est situé 14 rue de Londres, 75009 Paris, représentée par Madame Claire Fontaine en qualité de Présidente, ci-après dénommée « le Prestataire ».",
        "Et : Groupe Meridian SA, société anonyme au capital de 1 200 000 euros, immatriculée au RCS de Lyon sous le numéro 401 228 940, dont le siège social est situé 8 avenue Jean Jaurès, 69007 Lyon, représentée par Monsieur Paul Vasseur en qualité de Directeur Général, ci-après dénommée « le Client ».",
        "Il a été convenu ce qui suit.",
    ]),

    ("heading", "Article 1 — Objet du contrat"),
    ("p", [
        "Le présent contrat a pour objet de définir les conditions dans lesquelles le Prestataire réalise, pour le compte du Client, une mission d'accompagnement à la refonte de son système d'information de gestion.",
        "La mission comprend un audit de l'existant, la rédaction des spécifications fonctionnelles et l'assistance au déploiement.",
    ]),

    ("heading", "Article 2 — Documents contractuels"),
    ("p", [
        "Les documents contractuels comprennent le présent contrat et son annexe 1 décrivant le détail des prestations.",
        {
            "text": "Les livrables mentionnés à l'article 2 sera remis au Client selon le calendrier convenu entre les parties.",
            "defect": {
                "skill": "grammar",
                "suggestion": "Les livrables mentionnés à l'article 2 seront remis au Client selon le calendrier convenu entre les parties.",
                "explanation": "Accord sujet-verbe : le sujet « les livrables » est au pluriel, le verbe doit l'être aussi.",
                "priority": "high",
                "confidence": 0.97,
            },
        },
        "En cas de contradiction entre le contrat et son annexe, les stipulations du contrat prévalent.",
    ]),

    ("heading", "Article 3 — Durée"),
    ("p", [
        "Le présent contrat est conclu pour une durée de douze mois à compter du 1er mars 2025.",
        {
            "text": "Le présent contrat prend effet à la date de signature et court jusqu'à sa resiliation dans les conditions prévues à l'article 11.",
            "defect": {
                "skill": "spelling",
                "suggestion": "Le présent contrat prend effet à la date de signature et court jusqu'à sa résiliation dans les conditions prévues à l'article 11.",
                "explanation": "« resiliation » s'écrit « résiliation », avec un accent aigu.",
                "priority": "medium",
                "confidence": 0.95,
            },
        },
        {
            "text": "Il est expressément convenu que le contrat prend effet le 1er avril 2025.",
            "defect": {
                "skill": "consistency",
                "suggestion": "Il est expressément convenu que le contrat prend effet le 1er mars 2025.",
                "explanation": "Contradiction avec la date d'entrée en vigueur fixée en début d'article 3 (1er mars 2025).",
                "priority": "high",
                "confidence": 0.93,
            },
        },
        {
            "text": "Chacune des parties s'engagent à respecter le calendrier figurant en annexe.",
            "defect": {
                "skill": "grammar",
                "suggestion": "Chacune des parties s'engage à respecter le calendrier figurant en annexe.",
                "explanation": "« Chacune » est un pronom singulier : le verbe s'accorde au singulier.",
                "priority": "medium",
                "confidence": 0.94,
            },
        },
    ]),

    ("heading", "Article 4 — Conditions financières"),
    ("p", [
        {
            "text": "Le montant total des prestations est fixé à 45 000 euros hors taxes.",
            "defect": {
                "skill": "consistency",
                "suggestion": "Le montant total des prestations est fixé à 48 000 euros hors taxes.",
                "explanation": "Le budget indiqué en annexe 1 est de 48 000 euros hors taxes : les deux montants doivent concorder.",
                "priority": "high",
                "confidence": 0.91,
            },
        },
        "Ce montant est ferme et non révisable pendant toute la durée du contrat.",
        "Les frais de déplacement engagés par le Prestataire sont refacturés au réel, sur présentation des justificatifs.",
    ]),

    ("heading", "Article 5 — Modalités de paiement"),
    ("p", [
        "Les factures sont émises mensuellement et payables à trente jours fin de mois.",
        {
            "text": "Tout paiment intervenant après cette date donnera lieu à des pénalités de retard calculées au taux légal.",
            "defect": {
                "skill": "spelling",
                "suggestion": "Tout paiement intervenant après cette date donnera lieu à des pénalités de retard calculées au taux légal.",
                "explanation": "« paiment » s'écrit « paiement ».",
                "priority": "medium",
                "confidence": 0.96,
            },
        },
    ]),

    ("heading", "Article 6 — Obligations du Prestataire"),
    ("p", [
        {
            "text": "Novatech Conseil met en œuvre les moyens humains et techniques nécessaires à la bonne exécution de la mission.",
            "defect": {
                "skill": "consistency",
                "suggestion": "NOVATEK Conseil met en œuvre les moyens humains et techniques nécessaires à la bonne exécution de la mission.",
                "explanation": "Le Prestataire est désigné « NOVATEK Conseil » dans le préambule : la dénomination sociale doit être identique partout.",
                "priority": "high",
                "confidence": 0.9,
            },
        },
        {
            "text": "Le Prestataire remettra une PSSI conforme aux exigences du Client.",
            "defect": {
                "skill": "consistency",
                "suggestion": "Le Prestataire remettra une politique de sécurité des systèmes d'information (PSSI) conforme aux exigences du Client.",
                "explanation": "Le sigle PSSI est employé ici avant d'être défini, ce qui n'intervient qu'à l'article 8.",
                "priority": "medium",
                "confidence": 0.86,
            },
        },
        {
            "text": "On fera au mieux pour livrer dans les temps.",
            "defect": {
                "skill": "tone",
                "suggestion": "Le Prestataire met en œuvre les moyens nécessaires au respect des délais convenus.",
                "explanation": "Registre familier et engagement imprécis, incompatibles avec la portée juridique d'un contrat.",
                "priority": "high",
                "confidence": 0.92,
            },
        },
    ]),

    ("heading", "Article 7 — Obligations du Client"),
    ("p", [
        "Le Client met à disposition du Prestataire les informations et accès nécessaires à la réalisation de la mission.",
        {
            "text": "Le Client devra juste nous prévenir en cas de souci sur les accès.",
            "defect": {
                "skill": "tone",
                "suggestion": "Le Client informe le Prestataire sans délai de toute difficulté affectant la mise à disposition des accès.",
                "explanation": "Formulation familière et first-person inadaptée : un contrat désigne les parties par leur qualité.",
                "priority": "high",
                "confidence": 0.93,
            },
        },
        "Le Client désigne un interlocuteur unique chargé du suivi de la mission.",
    ]),

    ("heading", "Article 8 — Confidentialité"),
    ("p", [
        "Les parties s'engagent à préserver la confidentialité des informations échangées dans le cadre du présent contrat, et notamment la politique de sécurité des systèmes d'information (PSSI) du Client.",
        {
            "text": "Le Prestataire garantit la confidentailité des informations transmises par le Client pendant toute la durée du contrat.",
            "defect": {
                "skill": "spelling",
                "suggestion": "Le Prestataire garantit la confidentialité des informations transmises par le Client pendant toute la durée du contrat.",
                "explanation": "« confidentailité » s'écrit « confidentialité ».",
                "priority": "medium",
                "confidence": 0.95,
            },
        },
        {
            "text": "Cette obligation demeure en vigueur pendant cinq années à compter du terme du contrat, sauf si les informations concernées sont tombées dans le domaine public sans qu'aucune faute ne soit imputable à la partie qui les a reçues, étant précisé que la charge de la preuve de cette divulgation antérieure, qui doit être établie par tout moyen écrit, incombe à celle-ci, laquelle devra en informer l'autre partie dans un délai raisonnable.",
            "defect": {
                "skill": "clarity",
                "suggestion": "Cette obligation demeure en vigueur pendant cinq années à compter du terme du contrat. Elle cesse si les informations sont tombées dans le domaine public sans faute de la partie qui les a reçues. Il appartient à cette partie d'en apporter la preuve écrite et d'en informer l'autre sans délai.",
                "explanation": "Phrase de plus de soixante mots enchaînant quatre subordonnées : la découper rend l'obligation lisible.",
                "priority": "medium",
                "confidence": 0.88,
            },
        },
    ]),

    ("heading", "Article 9 — Propriété intellectuelle"),
    ("p", [
        "Les livrables réalisés dans le cadre de la mission deviennent la propriété du Client à compter de leur paiement intégral.",
        {
            "text": "Le Prestataire conserve la propriété de ses méthodes et outils préexistants, celui-ci ne pouvant les revendiquer.",
            "defect": {
                "skill": "clarity",
                "suggestion": "Le Prestataire conserve la propriété de ses méthodes et outils préexistants, que le Client ne peut revendiquer.",
                "explanation": "« celui-ci » peut désigner le Prestataire comme le Client : l'ambiguïté porte sur qui ne peut rien revendiquer.",
                "priority": "high",
                "confidence": 0.89,
            },
        },
    ]),

    ("heading", "Article 10 — Responsabilité"),
    ("p", [
        "La responsabilité du Prestataire est limitée aux dommages directs et plafonnée au montant total des sommes effectivement versées au titre du présent contrat.",
        "Le Prestataire justifie d'une assurance responsabilité civile professionnelle en cours de validité.",
    ]),

    ("heading", "Article 11 — Résiliation"),
    ("p", [
        "Chaque partie peut résilier le présent contrat par lettre recommandée avec accusé de réception, moyennant un préavis de deux mois.",
        "En cas de manquement grave de l'une des parties à ses obligations, l'autre partie peut résilier le contrat de plein droit quinze jours après une mise en demeure restée sans effet.",
    ]),

    ("heading", "Article 12 — Droit applicable et litiges"),
    ("p", [
        "Le présent contrat est soumis au droit français.",
        "À défaut de résolution amiable, tout litige relatif à sa validité, son interprétation ou son exécution relève de la compétence exclusive du Tribunal de commerce de Paris.",
        "Fait à Paris, en deux exemplaires originaux, le 24 février 2025.",
    ]),

    ("heading", "Annexe 1 — Description des prestations"),
    ("p", [
        "La mission se déroule en trois phases : cadrage, spécifications et accompagnement au déploiement.",
        "La phase de cadrage comprend l'audit de l'existant et la restitution d'un diagnostic écrit.",
        "Le budget global de la mission s'élève à 48 000 euros hors taxes, réparti entre les trois phases.",
        "Le calendrier prévisionnel prévoit une restitution finale au plus tard le 28 février 2026.",
    ]),
]


def sentences_of(paragraph):
    """Yields (text, defect_or_None) for every sentence of a paragraph."""
    for item in paragraph:
        if isinstance(item, str):
            yield item, None
        else:
            yield item["text"], item["defect"]


def build_docx(path):
    from docx import Document
    from docx.shared import Pt

    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for kind, content in CONTRACT:
        if kind == "heading":
            level = 0 if content.startswith("CONTRAT") else 1
            document.add_heading(content, level)
        else:
            # One paragraph holding all its sentences, as in a real contract.
            document.add_paragraph(" ".join(text for text, _ in sentences_of(content)))

    os.makedirs(os.path.dirname(path), exist_ok=True)
    document.save(path)
    return path


# The app paginates DOCX by character budget; mirroring it keeps the demo
# preview and the imported file structurally comparable.
MAX_CHARS_PER_PAGE = 3500


def build_mocks(document_path, findings_path):
    blocks = []
    for kind, content in CONTRACT:
        if kind == "heading":
            blocks.append({"kind": "heading", "text": content, "defect": None})
        else:
            for text, defect in sentences_of(content):
                blocks.append({"kind": "p", "text": text, "defect": defect})

    pages, current, size = [], [], 0
    for block in blocks:
        if size + len(block["text"]) > MAX_CHARS_PER_PAGE and current:
            pages.append(current)
            current, size = [], 0
        current.append(block)
        size += len(block["text"])
    if current:
        pages.append(current)

    doc_lines = [
        "// GENERATED — do not edit by hand.",
        "// Source: tools/build-sample-contract.py, the same one that produces",
        "// public/samples/contrat-prestation-services.docx. The demo engine matches",
        "// findings to sentences by exact text, so both must come from one source.",
        "//",
        "// Run: python3 tools/build-sample-contract.py",
        "",
        "export const MOCK_DOCUMENT = {",
        '  kind: "docx",',
        '  title: "Contrat de prestation de services.docx",',
        '  subtitle: "NOVATEK Conseil SAS — Groupe Meridian SA",',
        "  pages: [",
    ]
    findings = []

    for page_index, page in enumerate(pages, start=1):
        doc_lines.append(f"    // ── Page {page_index} ─────────────────────────────")
        doc_lines.append("    [")
        for block in page:
            marker = f"  // FINDING {block['defect']['skill']}" if block["defect"] else ""
            doc_lines.append(
                f"      {{ kind: {json.dumps(block['kind'])}, text: {json.dumps(block['text'], ensure_ascii=False)} }},{marker}"
            )
            if block["defect"]:
                findings.append({
                    "skill": block["defect"]["skill"],
                    "page": page_index,
                    "original": block["text"],
                    "suggestion": block["defect"]["suggestion"],
                    "explanation": block["defect"]["explanation"],
                    "priority": block["defect"]["priority"],
                    "confidence": block["defect"]["confidence"],
                })
        doc_lines.append("    ],")

    doc_lines += ["  ],", "};", ""]
    with open(document_path, "w", encoding="utf-8") as handle:
        handle.write("\n".join(doc_lines))

    findings.sort(key=lambda f: (f["skill"], f["page"]))
    finding_lines = [
        "// GENERATED — do not edit by hand.",
        "// Source: tools/build-sample-contract.py",
        "//",
        "// Every `original` matches a sentence of mockDocument.js exactly (same page,",
        "// same text), which is what lets the preview highlight it in demo mode.",
        "",
        "export const MOCK_FINDINGS_POOL = [",
    ]
    for finding in findings:
        finding_lines.append("  {")
        for key in ("skill", "page", "original", "suggestion", "explanation", "priority", "confidence"):
            value = finding[key]
            rendered = value if isinstance(value, (int, float)) else json.dumps(value, ensure_ascii=False)
            finding_lines.append(f"    {key}: {rendered},")
        finding_lines.append("  },")
    finding_lines += ["];", ""]

    # Custom checks keep their existing shape: they are not tied to a sentence
    # of the document, so the app can attach them anywhere.
    finding_lines += [
        "/** Templates used to fabricate a finding for each user-defined check. */",
        "export const CUSTOM_CHECK_TEMPLATES = [",
        "  {",
        "    page: 2,",
        '    original: "Le montant total des prestations est fixé à 45 000 euros hors taxes.",',
        '    suggestion: "Le montant total des prestations est fixé à 48 000 euros hors taxes (cf. annexe 1).",',
        '    explanationTpl: "Contrôle « {name} » : le montant ne correspond pas à celui de l\'annexe.",',
        '    priority: "high",',
        "    confidence: 0.9,",
        "  },",
        "  {",
        "    page: 1,",
        '    original: "Le présent contrat est conclu pour une durée de douze mois à compter du 1er mars 2025.",',
        '    suggestion: "Le présent contrat est conclu pour une durée de douze mois à compter du 1er mars 2025, renouvelable par tacite reconduction.",',
        '    explanationTpl: "Contrôle « {name} » : la clause de reconduction attendue est absente.",',
        '    priority: "medium",',
        "    confidence: 0.84,",
        "  },",
        "  {",
        "    page: 3,",
        '    original: "Le présent contrat est soumis au droit français.",',
        '    suggestion: "Le présent contrat est soumis au droit français, à l\'exclusion de ses règles de conflit de lois.",',
        '    explanationTpl: "Contrôle « {name} » : la clause de droit applicable gagnerait à être précisée.",',
        '    priority: "low",',
        "    confidence: 0.78,",
        "  },",
        "];",
        "",
    ]
    with open(findings_path, "w", encoding="utf-8") as handle:
        handle.write("\n".join(finding_lines))

    return len(pages), len(findings)


if __name__ == "__main__":
    docx_path = build_docx("public/samples/contrat-prestation-services.docx")
    pages, defects = build_mocks("src/data/mockDocument.js", "src/data/mockFindings.js")
    print(f"{docx_path}")
    print(f"src/data/mockDocument.js — {pages} pages")
    print(f"src/data/mockFindings.js — {defects} defauts plantes")
